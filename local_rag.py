import torch
import ollama
import os
import json
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import argparse

# ANSI escape codes for colors
PINK = '\033[95m'
CYAN = '\033[96m'
YELLOW = '\033[93m'
NEON_GREEN = '\033[92m'
RESET_COLOR = '\033[0m'

# Function to open a file and return its contents as a string
def open_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as infile:
        return infile.read()

# Function to get relevant context from the vault based on user input
def get_relevant_context(rewritten_input, vault_embeddings, vault_content, top_k=3):
    if vault_embeddings.nelement() == 0:  # Check if the tensor has any elements
        return []
    # Encode the rewritten input
    input_embedding = ollama.embeddings(model='mxbai-embed-large', prompt=rewritten_input)["embedding"]
    # Compute cosine similarity between the input and vault embeddings
    cos_scores = torch.cosine_similarity(torch.tensor(input_embedding).unsqueeze(0), vault_embeddings)
    # Adjust top_k if it's greater than the number of available scores
    top_k = min(top_k, len(cos_scores))
    # Sort the scores and get the top-k indices
    top_indices = torch.topk(cos_scores, k=top_k)[1].tolist()
    # Get the corresponding context from the vault
    relevant_context = [vault_content[idx].strip() for idx in top_indices]
    return relevant_context

def rewrite_query(user_input_json, conversation_history, ollama_model):
    user_input = json.loads(user_input_json)["Query"]
    context = "\n".join([f"{msg['role']}: {msg['content']}" for msg in conversation_history[-2:]])
    prompt = f"""Rewrite the following query by incorporating relevant context from the conversation history.
    The rewritten query should:
    
    - Preserve the core intent and meaning of the original query
    - Expand and clarify the query to make it more specific and informative for retrieving relevant context
    - Give answers based on the text you were trained with.
    - Give the answers only in romanian instead of English.
    - You will be focusing only on the following aspects:
        1. Technologies Used: Specify the technologies that will be used for the project, including any frameworks, tools, or platforms.
        2. Proposal Structure: Provide a detailed structure of the proposal, including the main sections and their contents.
        3.Additional Suggestions: Offer any additional suggestions that might enhance the proposal or address potential improvements.
        4.Price and Implementation Time: Include an estimated price and time required for implementation.
    
    Return ONLY the rewritten query text, without any additional formatting or explanations.
    
    Conversation History:
    {context}
    
    Original query: [{user_input}]
    
    Rewritten query: 
    """
    response = client.chat.completions.create(
        model=ollama_model,
        messages=[{"role": "system", "content": prompt}],
        max_tokens=200,
        n=1,
        temperature=0.1,
    )
    rewritten_query = response.choices[0].message.content.strip()
    return json.dumps({"Rewritten Query": rewritten_query})
   
def ollama_chat(user_input, system_message, vault_embeddings, vault_content, ollama_model, conversation_history):
    conversation_history.append({"role": "user", "content": user_input})
    
    if len(conversation_history) > 1:
        query_json = {
            "Query": user_input,
            "Rewritten Query": ""
        }
        rewritten_query_json = rewrite_query(json.dumps(query_json), conversation_history, ollama_model)
        rewritten_query_data = json.loads(rewritten_query_json)
        rewritten_query = rewritten_query_data["Rewritten Query"]
        print(PINK + "Original Query: " + user_input + RESET_COLOR)
        print(PINK + "Rewritten Query: " + rewritten_query + RESET_COLOR)
    else:
        rewritten_query = user_input
    
    relevant_context = get_relevant_context(rewritten_query, vault_embeddings, vault_content)
    if relevant_context:
        context_str = "\n".join(relevant_context)
        print("Context Pulled from Documents: \n\n" + CYAN + context_str + RESET_COLOR)
    else:
        print(CYAN + "No relevant context found." + RESET_COLOR)
    
    user_input_with_context = user_input
    if relevant_context:
        user_input_with_context = user_input + "\n\nRelevant Context:\n" + context_str
    
    conversation_history[-1]["content"] = user_input_with_context
    
    messages = [
        {"role": "system", "content": system_message},
        *conversation_history
    ]
    
    response = client.chat.completions.create(
        model=ollama_model,
        messages=messages,
        max_tokens=2000,
    )
    
    conversation_history.append({"role": "assistant", "content": response.choices[0].message.content})
    
    return response.choices[0].message.content

def save_to_docx(request_text, response_text, filename):
    doc = Document()
    
    # Set document title
    title = doc.add_heading(level=1)
    run = title.add_run("Solicitarea client:")
    run.font.bold = True
    run.font.size = Pt(16)
    
    # Add client request
    request_paragraph = doc.add_paragraph(request_text)
    request_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Add a separator line
    separator = doc.add_paragraph()
    separator_run = separator.add_run()
    separator_run.add_break()
    separator_run.add_break()
    
    # Set scopul documentului title
    scopul_title = doc.add_heading(level=1)
    run = scopul_title.add_run("I. Scopul documentului:")
    run.font.bold = True
    run.font.size = Pt(16)
    
    scopul_text = """
Această ofertă preliminară este bazată pe informațiile pe care ni le-ați furnizat. Înainte de a începe efectiv dezvoltarea și de a vă putea oferi o estimare exactă a costurilor și timpului necesar, va fi necesar să parcurgem câteva etape esențiale de planificare

1. Elaborarea unei diagrame logice pentru a defini arhitectura aplicației. 
2. Crearea unei diagrame ER pentru a structura baza de date. 
3. Realizarea unui design inițial în Figma pentru a elimina orice ambiguitate legată de interfața grafică. 

Pentru începerea etapelor de mai sus, va fi necesar să semnăm un contract de colaborare și să achitați în avans prețul acestora: 

Diagrama logică și diagrama ER: X Euro + TVA 
Proiectul în Figma: X Euro + TVA 
Acești bani se vor scădea din prețul total de dezvoltare odată acceptată oferta fermă. 

Oferta de mai jos este doar orientativă și urmărește să vă ofere o perspectivă asupra modului în care operăm și a costurilor probabile. Pentru o ofertă și un timp de implementare exact, va fi necesar să completăm etapele de planificare menționate.
"""
    
    # Add scopul text
    scopul_paragraph = doc.add_paragraph(scopul_text)
    scopul_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Add a separator line
    separator = doc.add_paragraph()
    separator_run = separator.add_run()
    separator_run.add_break()
    separator_run.add_break()
    
    # Set document offer title
    offer_title = doc.add_heading(level=1)
    run = offer_title.add_run("Oferta pentru firmă:")
    run.font.bold = True
    run.font.size = Pt(16)
    
    # Add the offer content
    response_paragraph = doc.add_paragraph(response_text)
    response_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    doc.save(filename)


# Parse command-line arguments
print(NEON_GREEN + "Parsing command-line arguments..." + RESET_COLOR)
parser = argparse.ArgumentParser(description="Ollama Chat")
parser.add_argument("--model", default="llama3", help="Ollama model to use (default: llama3)")
args = parser.parse_args()

# Configuration for the Ollama API client
print(NEON_GREEN + "Initializing Ollama API client..." + RESET_COLOR)
client = OpenAI(
    base_url='http://localhost:11434/v1',
    api_key='llama3'
)

# Load the vault content
print(NEON_GREEN + "Loading vault content..." + RESET_COLOR)
vault_content = []
if os.path.exists("vault.txt"):
    with open("vault.txt", "r", encoding='utf-8') as vault_file:
        vault_content = vault_file.readlines()

# Generate embeddings for the vault content using Ollama
print(NEON_GREEN + "Generating embeddings for the vault content..." + RESET_COLOR)
vault_embeddings = []
for content in vault_content:
    response = ollama.embeddings(model='mxbai-embed-large', prompt=content)
    vault_embeddings.append(response["embedding"])

# Convert to tensor and print embeddings
print("Converting embeddings to tensor...")
vault_embeddings_tensor = torch.tensor(vault_embeddings) 
print("Embeddings for each line in the vault:")
print(vault_embeddings_tensor) 
print("Embedding tensor dimensions:", vault_embeddings_tensor.shape)

# Initialize conversation history
conversation_history = []

# Main loop
print(YELLOW + "Starting main loop. Type your message and press Enter to interact with the model." + RESET_COLOR)
print(YELLOW + "Type 'exit' to quit." + RESET_COLOR)

while True:
    user_input = input(PINK + "You: " + RESET_COLOR)
    
    if user_input.lower() == 'exit':
        print(NEON_GREEN + "Exiting the program. Goodbye!" + RESET_COLOR)
        break
    
    # Ollama model and system message
    ollama_model = args.model
    system_message = "You are a helpful assistant. Please respond only in romanian. In your answers, ou will have in cosideration the technologies that will be used, you will give a layout proposal with the different apps that will need to be implemented and also you will offer informations about the time an price for the full implemenations. "
    
    # Call the ollama_chat function
    response = ollama_chat(user_input, system_message, vault_embeddings_tensor, vault_content, ollama_model, conversation_history)
    
    # Print the assistant's response
    print(YELLOW + "Assistant: " + RESET_COLOR + response)
    
    # Save the conversation to a DOCX file
    request_text = "Solicitarea clientului: " + user_input
    response_text = response
    filename = "solicitarea_si_oferta.docx"
    save_to_docx(request_text, response_text, filename)
    print(NEON_GREEN + "Document saved as " + filename + RESET_COLOR)
